#!/usr/bin/env python3
"""
Dual System Evaluation Word Document Generator

This script generates Microsoft Word documents for minister evaluation by:
1. Loading retrieval results from Step 2 (dual_system_retrieval.py output)
2. Creating Word documents with 10 documents per file (5 from each system)
3. Blinded presentation (no system identification)
4. Judge field and instructions at the top
5. Sequential numbering for easy processing

Usage:
    python generate_word_docs.py --results 3large_vs_3small/step2_retrieval_results.json --output-dir 3large_vs_3small/word_docs
"""

import argparse
import json
import os
import random
import re
import sys
from typing import Any

from docx import Document
from docx.shared import Inches, Pt


def sanitize_filename(filename: str) -> str:
    """Convert string to valid filename."""
    sanitized = re.sub(r"[^\w\s-]", "_", filename)
    sanitized = re.sub(r"[\s]+", "-", sanitized)
    if len(sanitized) > 50:
        sanitized = sanitized[:50]
    return sanitized


def create_word_document(
    query: str, documents: list[dict[str, Any]], output_path: str
) -> None:
    """Create a Word document for minister evaluation."""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add Judge field
    judge_para = doc.add_paragraph()
    judge_para.add_run("Judge: ").bold = True

    # Add spacing
    doc.add_paragraph()

    # Add Query with prominent formatting
    doc.add_heading("Query", level=1)

    # Create a bordered paragraph for the query
    query_para = doc.add_paragraph()
    query_run = query_para.add_run(query)
    query_run.bold = True
    query_run.font.size = Pt(14)  # Larger font size

    # Add border around the query paragraph
    from docx.oxml.shared import OxmlElement, qn

    # Add border to the query paragraph
    pPr = query_para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pBdr.set(qn("w:top"), "single")
    pBdr.set(qn("w:left"), "single")
    pBdr.set(qn("w:bottom"), "single")
    pBdr.set(qn("w:right"), "single")
    pBdr.set(qn("w:space"), "4")
    pBdr.set(qn("w:sz"), "12")
    pPr.append(pBdr)

    # Add Instructions
    doc.add_heading("Instructions", level=2)
    doc.add_paragraph("Review each document below and assign a relevance score:")

    # Add scoring scale
    scale_items = [
        "3: Highly Relevant - Directly answers the query",
        "2: Relevant - Contains information related to the query",
        "1: Marginally Relevant - Mentions query topics but not directly helpful",
        "0: Irrelevant - Not related to the query",
        "ignore: Skip this document (write 'I' or 'ignore', or leave blank)",
    ]

    for item in scale_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()

    # Add Documents section
    doc.add_heading("Documents", level=2)

    # Add each document
    for i, document in enumerate(documents, 1):
        # Document header
        doc.add_heading(f"Document {i}", level=3)

        # Document content
        doc.add_paragraph(document["text"])

        # Metadata (if available and useful)
        if document.get("metadata"):
            metadata = document["metadata"]
            filtered_metadata = {
                k: v
                for k, v in metadata.items()
                if k in ["title", "author", "source", "type"] and v
            }
            if filtered_metadata:
                doc.add_paragraph()
                metadata_para = doc.add_paragraph("Metadata:")
                metadata_para.italic = True
                for key, value in filtered_metadata.items():
                    doc.add_paragraph(f"{key.title()}: {value}", style="List Bullet")

        # Scoring section
        doc.add_paragraph()
        scoring_para = doc.add_paragraph()
        scoring_para.add_run("Scoring: ").italic = True
        scoring_para.add_run(
            "0=Irrelevant, 1=Marginally Relevant, 2=Relevant, 3=Highly Relevant, ignore=Skip"
        ).italic = True

        score_para = doc.add_paragraph()
        score_para.add_run("Relevance Score [Enter 0-3 or ignore]: ").bold = True

        # Add separator line
        if i < len(documents):
            doc.add_paragraph("-" * 80)
            doc.add_paragraph()

    # Save document
    doc.save(output_path)


def load_retrieval_results(results_file: str) -> dict[str, Any]:
    """Load retrieval results from Step 2."""
    try:
        with open(results_file) as f:
            return json.load(f)
    except FileNotFoundError:
        sys.exit(f"Error: Results file '{results_file}' not found.")
    except json.JSONDecodeError as e:
        sys.exit(f"Error: Invalid JSON in results file: {e}")


def prepare_documents_for_query(query_data: dict[str, Any]) -> list[dict[str, Any]]:
    """Prepare and randomize documents from both systems for a query."""
    documents = []

    # Handle dual_system_retrieval.py format
    systems_data = query_data.get("systems", {})

    for system_name, system_info in systems_data.items():
        system_docs = system_info.get("documents", [])

        # Add system tracking (for later processing) but don't show in Word doc
        for doc in system_docs:
            doc_copy = doc.copy()
            doc_copy["_system"] = system_name  # Hidden field for processing
            documents.append(doc_copy)

    # Randomize order for blinded evaluation
    random.shuffle(documents)

    return documents


def generate_word_documents(results_file: str, output_dir: str) -> None:
    """Generate Word documents for all queries in the results."""
    # Create directory structure
    unassigned_dir = os.path.join(output_dir, "unassigned")
    todo_dir = os.path.join(output_dir, "todo")
    done_dir = os.path.join(output_dir, "done")

    for directory in [unassigned_dir, todo_dir, done_dir]:
        os.makedirs(directory, exist_ok=True)

    # Load retrieval results
    results = load_retrieval_results(results_file)

    # Handle dual_system_retrieval.py format
    if isinstance(results, dict) and "results" in results:
        queries_list = results["results"]
    else:
        # Legacy format - assume it's a direct list or individual queries
        queries_list = results if isinstance(results, list) else [results]

    # Process each query
    queries_processed = 0

    for query_data in queries_list:
        query_id = query_data.get("query_id", "unknown")
        query_text = query_data.get("query_text", f"Query {query_id}")

        print(f"Processing query: {query_text}")

        # Prepare documents (5 from each system, randomized)
        documents = prepare_documents_for_query(query_data)

        if len(documents) != 10:
            print(
                f"Warning: Expected 10 documents for query '{query_text}', got {len(documents)}"
            )
            continue

        # Create filename
        filename = f"{sanitize_filename(query_text)}.docx"
        filepath = os.path.join(unassigned_dir, filename)

        # Handle filename conflicts
        counter = 1
        while os.path.exists(filepath):
            base_name = f"{sanitize_filename(query_text)}_{counter}"
            filename = f"{base_name}.docx"
            filepath = os.path.join(unassigned_dir, filename)
            counter += 1

        # Create Word document
        try:
            create_word_document(query_text, documents, filepath)
            print(f"Created: {filename}")
            queries_processed += 1
        except Exception as e:
            print(f"Error creating document for query '{query_text}': {e}")
            continue

    print("\nGeneration complete!")
    print(f"Processed {queries_processed} queries")
    print(f"Word documents created in: {unassigned_dir}")
    print("\nNext steps:")
    print(f"1. Distribute documents from {unassigned_dir} to ministers")
    print(f"2. When completed, place them in {todo_dir}")
    print("3. Run process_word_docs.py to extract judgments")


def main():
    parser = argparse.ArgumentParser(
        description="Generate Word documents for dual system evaluation"
    )
    parser.add_argument(
        "--results",
        type=str,
        required=True,
        help="Path to Step 2 retrieval results JSON file",
    )
    parser.add_argument(
        "--output-dir",
        type=str,
        required=True,
        help="Output directory for Word documents",
    )
    parser.add_argument(
        "--seed",
        type=int,
        default=42,
        help="Random seed for document shuffling (default: 42)",
    )

    args = parser.parse_args()

    # Set random seed for reproducible shuffling
    random.seed(args.seed)

    generate_word_documents(args.results, args.output_dir)


if __name__ == "__main__":
    main()
